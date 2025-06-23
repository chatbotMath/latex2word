import streamlit as st
import re
import os
import subprocess
import tempfile
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io

class LaTeXToWordConverter:
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        
    def extract_exercises(self, latex_content):
        """Extract all exercises from LaTeX content"""
        # Pattern to match \begin{ex} ... \end{ex}
        pattern = r'\\begin\{ex\}(.*?)\\end\{ex\}'
        exercises = re.findall(pattern, latex_content, re.DOTALL)
        return exercises
    
    def parse_exercise(self, exercise_content):
        """Parse individual exercise content"""
        # Extract question (content before \choice or \immini)
        question_match = re.search(r'\\immini\{(.*?)\}', exercise_content, re.DOTALL)
        if not question_match:
            # Try to find question before \choice
            parts = exercise_content.split('\\choice')
            question = parts[0].strip() if parts else ""
        else:
            question = question_match.group(1).strip()
        
        # Extract choices and find correct answer
        choice_pattern = r'\{([^}]*)\}'
        choices_section = re.search(r'\\choice(.*?)(?=\\begin\{tikzpicture\}|\\begin\{tabular\}|\\loigiai|$)', exercise_content, re.DOTALL)
        choices = []
        correct_choice_index = -1
        
        if choices_section:
            raw_choices = re.findall(choice_pattern, choices_section.group(1))
            for idx, choice in enumerate(raw_choices):
                choice = choice.strip()
                if choice:
                    # Check if this is the correct answer
                    if choice.startswith('\\True'):
                        correct_choice_index = idx
                        # Remove \True marker and clean the choice
                        choice = choice.replace('\\True', '').strip()
                    choices.append(choice)
        
        # Extract TikZ picture
        tikz_match = re.search(r'\\begin\{tikzpicture\}(.*?)\\end\{tikzpicture\}', exercise_content, re.DOTALL)
        tikz_content = tikz_match.group(0) if tikz_match else None
        
        # Extract tables
        tables = self.extract_and_convert_tables(exercise_content)
        
        # Extract solution
        solution_match = re.search(r'\\loigiai\{(.*?)\}', exercise_content, re.DOTALL)
        solution = solution_match.group(1).strip() if solution_match else None
        
        return {
            'question': question,
            'choices': choices,
            'correct_choice': correct_choice_index,
            'tikz': tikz_content,
            'tables': tables,
            'solution': solution
        }
    
    def extract_and_convert_tables(self, content):
        """Extract LaTeX tables and convert to markdown format"""
        tables = []
        
        # Find all tabular environments
        tabular_pattern = r'\\begin\{tabular\}(\{[^}]*\})(.*?)\\end\{tabular\}'
        tabular_matches = re.finditer(tabular_pattern, content, re.DOTALL)
        
        for match in tabular_matches:
            column_spec = match.group(1)
            table_content = match.group(2)
            
            # Parse column specification
            col_count = len(re.findall(r'[lcr]', column_spec))
            
            # Convert table to markdown
            markdown_table = self.latex_table_to_markdown(table_content, col_count)
            tables.append(markdown_table)
            
        return tables
    
    def latex_table_to_markdown(self, table_content, col_count):
        """Convert LaTeX table content to markdown table"""
        # Clean up the content
        lines = table_content.strip().split('\\\\')
        lines = [line.strip() for line in lines if line.strip()]
        
        # Remove \hline commands
        lines = [line for line in lines if not line.strip() == '\\hline']
        
        markdown_rows = []
        
        for line in lines:
            # Split by & and clean each cell
            cells = line.split('&')
            cells = [self.clean_latex_text(cell.strip()) for cell in cells]
            
            # Pad row if necessary
            while len(cells) < col_count:
                cells.append('')
            
            # Create markdown row
            markdown_row = '| ' + ' | '.join(cells) + ' |'
            markdown_rows.append(markdown_row)
        
        # Add header separator (assuming first row is header)
        if markdown_rows:
            separator = '| ' + ' | '.join(['---'] * col_count) + ' |'
            markdown_rows.insert(1, separator)
        
        return '\n'.join(markdown_rows)
    
    def create_table_from_markdown(self, doc, markdown_table):
        """Create a Word table from markdown format"""
        lines = markdown_table.strip().split('\n')
        if len(lines) < 3:  # Need at least header, separator, and one data row
            return
        
        # Parse the table
        rows = []
        for line in lines:
            if '---' in line:  # Skip separator line
                continue
            # Extract cells between pipes
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
        
        if not rows:
            return
        
        # Create table in Word
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        
        # Populate table
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                table.cell(i, j).text = cell_data
                # Bold the header row
                if i == 0:
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    def clean_latex_text(self, text):
        """Clean LaTeX commands from text for Word"""
        # Dòng này tìm kiếm văn bản được bao trong dấu '$' (ví dụ: $E=mc^2$)
        # và thay thế nó chỉ bằng nội dung bên trong (ví dụ: E=mc^2).
        text = re.sub(r'\$([^$]+)\$', r'\1', text)
        
        # Bạn có thể thêm các lệnh re.sub khác ở đây để xóa các lệnh LaTeX khác
        # Ví dụ: xóa \textit{...}
        # text = re.sub(r'\\textit{([^}]+)}', r'\1', text)
        
        return text
    
    def compile_tikz_to_image(self, tikz_code, filename_base):
        """Compile TikZ code to image"""
        # Create complete LaTeX document
        latex_doc = f"""
\\documentclass[border=5pt]{{standalone}}
\\usepackage{{tikz}}
\\usepackage{{amsmath}}
\\usepackage{{amssymb}}
\\begin{{document}}
{tikz_code}
\\end{{document}}
"""
        
        # Write to temporary file
        tex_file = os.path.join(self.temp_dir, f"{filename_base}.tex")
        with open(tex_file, 'w', encoding='utf-8') as f:
            f.write(latex_doc)
        
        try:
            # Compile with pdflatex
            subprocess.run(['pdflatex', '-interaction=nonstopmode', '-output-directory', self.temp_dir, tex_file], 
                         capture_output=True, check=True)
            
            # Convert PDF to PNG
            pdf_file = os.path.join(self.temp_dir, f"{filename_base}.pdf")
            png_file = os.path.join(self.temp_dir, f"{filename_base}.png")
            
            # Use pdftoppm for conversion (more reliable than Ghostscript)
            subprocess.run(['pdftoppm', '-png', '-r', '300', '-singlefile', pdf_file, 
                          os.path.join(self.temp_dir, filename_base)], 
                         capture_output=True, check=True)
            
            # Check if image was created
            if os.path.exists(png_file):
                return png_file
            else:
                # Try alternative conversion with ImageMagick
                subprocess.run(['convert', '-density', '300', pdf_file, png_file], 
                             capture_output=True, check=True)
                return png_file if os.path.exists(png_file) else None
                
        except subprocess.CalledProcessError as e:
            st.error(f"Error compiling TikZ: {e}")
            return None
        except Exception as e:
            st.error(f"Unexpected error: {e}")
            return None
    
    def add_underline_to_run(self, run):
        """Add underline formatting to a run"""
        run.font.underline = True
    
    def create_word_document(self, exercises):
        """Create Word document from parsed exercises"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Bài tập', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for idx, exercise in enumerate(exercises, 1):
            # Add question number and text
            question_para = doc.add_paragraph()
            question_para.add_run(f'Câu {idx}. ').bold = True
            question_para.add_run(self.clean_latex_text(exercise['question']))
            
            # Add choices with correct answer underlined
            for i, choice in enumerate(exercise['choices']):
                choice_para = doc.add_paragraph()
                choice_label = f'    {chr(65 + i)}. '  # A, B, C, D
                
                # Check if this is the correct answer
                if exercise.get('correct_choice', -1) == i:
                    # Add underlined choice label
                    choice_run = choice_para.add_run(choice_label)
                    self.add_underline_to_run(choice_run)
                    choice_run.bold = True
                    # Add choice text (also underlined)
                    text_run = choice_para.add_run(self.clean_latex_text(choice))
                    self.add_underline_to_run(text_run)
                else:
                    # Normal choice
                    choice_para.add_run(choice_label)
                    choice_para.add_run(self.clean_latex_text(choice))
            
            # Add tables if exist
            if exercise.get('tables'):
                for table_markdown in exercise['tables']:
                    doc.add_paragraph()  # Add spacing
                    self.create_table_from_markdown(doc, table_markdown)
                    doc.add_paragraph()  # Add spacing
            
            # Add TikZ image if exists
            if exercise['tikz']:
                image_file = self.compile_tikz_to_image(exercise['tikz'], f'tikz_{idx}')
                if image_file and os.path.exists(image_file):
                    doc.add_picture(image_file, width=Inches(3))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add solution if exists
            if exercise['solution']:
                # Add empty line before solution
                doc.add_paragraph()
                
                # Add "Lời giải:" on its own line
                solution_header = doc.add_paragraph()
                solution_header.add_run('Lời giải:').bold = True
                
                # Add solution content on new line
                solution_content = doc.add_paragraph()
                
                # Check if solution contains tables
                solution_tables = self.extract_and_convert_tables(exercise['solution'])
                if solution_tables:
                    # Clean solution text (remove table content)
                    clean_solution = re.sub(r'\\begin\{tabular\}.*?\\end\{tabular\}', '', exercise['solution'], flags=re.DOTALL)
                    solution_content.add_run(self.clean_latex_text(clean_solution))
                    
                    # Add solution tables
                    for table_markdown in solution_tables:
                        doc.add_paragraph()
                        self.create_table_from_markdown(doc, table_markdown)
                else:
                    solution_content.add_run(self.clean_latex_text(exercise['solution']))
            
            # Add spacing between exercises
            doc.add_paragraph()
            doc.add_paragraph()
        
        return doc
    
    def cleanup(self):
        """Clean up temporary files"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

def main():
    st.set_page_config(page_title="LaTeX to Word Converter", page_icon="📝")
    
    st.title("🔄 LaTeX to Word Converter")
    st.markdown("Chuyển đổi bài tập LaTeX sang định dạng Word")
    
    # Create two columns
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📥 Input LaTeX")
        
        # Text area for LaTeX input
        latex_input = st.text_area(
            "Nhập code LaTeX của bạn:",
            height=400,
            value=r"""\begin{ex}
\immini{Cho bảng số liệu sau về điểm thi của lớp:
\begin{tabular}{|c|c|c|c|}
\hline
\textbf{STT} & \textbf{Họ tên} & \textbf{Điểm} & \textbf{Xếp loại} \\
\hline
1 & Nguyễn Văn A & 8.5 & Giỏi \\
\hline
2 & Trần Thị B & 7.0 & Khá \\
\hline
3 & Lê Văn C & 9.0 & Giỏi \\
\hline
\end{tabular}

Hỏi có bao nhiêu học sinh đạt loại Giỏi?
\choice
{1 học sinh}
{\True 2 học sinh}
{3 học sinh}
{4 học sinh}
}
\loigiai{
Dựa vào bảng số liệu, ta thấy:
\begin{itemize}
\item Nguyễn Văn A: 8.5 điểm - Xếp loại Giỏi
\item Lê Văn C: 9.0 điểm - Xếp loại Giỏi
\end{itemize}
Vậy có 2 học sinh đạt loại Giỏi.

Đáp án đúng là B.
}
\end{ex}

\begin{ex}
\immini{Dựa vào hình vẽ (Hình b), hãy chọn khẳng định đúng trong các khẳng định sau đây?
\choice
{ Điểm $M$ nằm giữa $2$ điểm $N$ và $P$}
{\True Điểm $N$ nằm giữa $2$ điểm $M$ và $P$}
{ Điểm $P$ nằm giữa $2$ điểm $M$ và $N$}
{ Hai điểm $M$ và $P$ nằm cùng phía đối với điểm $N$}
}{\begin{tikzpicture}[scale=1]
\coordinate (M) at (0.5, 0);
\coordinate (N) at (2.5, 0);
\coordinate (P) at (4.5, 0);
\draw[thick] (0, 0) -- (5.5, 0);
\foreach \pt/\angle in {M/90, N/90, P/90} {
\draw[fill=white] (\pt) circle (1.5pt) +(\angle:3mm) node{$\pt$};
}
\node[below=5mm of N] {Hình $b$};
\end{tikzpicture}
}
\loigiai{
Theo hình vẽ, các điểm $M$, $N$, $P$ nằm trên một đường thẳng theo thứ tự từ trái sang phải.

Vậy đáp án đúng là B.
}
\end{ex}"""
        )
        
        # File upload option
        uploaded_file = st.file_uploader("Hoặc tải lên file .tex", type=['tex'])
        if uploaded_file:
            latex_input = uploaded_file.read().decode('utf-8')
    
    with col2:
        st.subheader("📤 Output")
        
        if st.button("🔄 Chuyển đổi sang Word", type="primary"):
            if latex_input:
                try:
                    with st.spinner("Đang xử lý..."):
                        # Create converter instance
                        converter = LaTeXToWordConverter()
                        
                        # Extract and parse exercises
                        exercises_raw = converter.extract_exercises(latex_input)
                        exercises_parsed = [converter.parse_exercise(ex) for ex in exercises_raw]
                        
                        # Create Word document
                        doc = converter.create_word_document(exercises_parsed)
                        
                        # Save to BytesIO
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        # Clean up
                        converter.cleanup()
                        
                        # Offer download
                        st.success("✅ Chuyển đổi thành công!")
                        st.download_button(
                            label="📥 Tải xuống file Word",
                            data=doc_io.getvalue(),
                            file_name="exercises_converted.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        # Show preview
                        st.info(f"Đã chuyển đổi {len(exercises_parsed)} câu hỏi")
                        
                except Exception as e:
                    st.error(f"❌ Lỗi: {str(e)}")
            else:
                st.warning("⚠️ Vui lòng nhập nội dung LaTeX")
    
    # Instructions
    with st.expander("📖 Hướng dẫn sử dụng"):
        st.markdown("""
        ### Cấu trúc LaTeX được hỗ trợ:
        
        1. **Câu hỏi**: Đặt trong `\\begin{ex}...\\end{ex}`
        2. **Nội dung câu hỏi**: Trong `\\immini{...}` hoặc trước `\\choice`
        3. **Các lựa chọn**: Sau `\\choice`, mỗi lựa chọn trong `{...}`
        4. **Đáp án đúng**: Đánh dấu bằng `{\\True ...}` - sẽ được gạch chân trong Word
        5. **Hình vẽ TikZ**: Trong `\\begin{tikzpicture}...\\end{tikzpicture}`
        6. **Bảng**: Trong `\\begin{tabular}{|c|c|...}...\\end{tabular}`
        7. **Lời giải**: Trong `\\loigiai{...}` - sẽ xuất hiện trên dòng riêng
        
        ### Tính năng đặc biệt:
        
        #### Đánh dấu đáp án đúng:
        ```latex
        \\choice
        {Phương án A}
        {\\True Phương án B}  ← Đáp án đúng, sẽ được gạch chân
        {Phương án C}
        {Phương án D}
        ```
        
        #### Format lời giải:
        - "Lời giải:" sẽ được in đậm trên dòng riêng
        - Nội dung lời giải xuất hiện ở dòng tiếp theo
        - Hỗ trợ bảng và danh sách trong lời giải
        
        ### Xử lý bảng LaTeX:
        - Bảng `tabular` được chuyển sang định dạng Markdown
        - Hàng đầu tiên tự động được làm đậm (header)
        - Hỗ trợ các định dạng `\\textbf{}`, `\\textit{}`, etc.
        - Bảng trong Word sẽ có đường viền và format chuẩn
        
        ### Lưu ý:
        - Ứng dụng sẽ tự động biên dịch hình TikZ thành ảnh
        - Các ký hiệu toán học sẽ được chuyển đổi sang text thường
        - Đáp án đúng được gạch chân và in đậm
        - File Word xuất ra sẽ có format chuẩn với đánh số câu hỏi tự động
        """)
    
    # Footer
    st.markdown("---")
    st.markdown("💡 **Tip**: Bạn có thể copy nhiều câu hỏi cùng lúc, mỗi câu trong `\\begin{ex}...\\end{ex}`")

if __name__ == "__main__":
    main()
, r'\1', text)  # Remove math delimiters
        text = re.sub(r'\\item', '•', text)
        text = re.sub(r'\\begin\{itemize\}', '', text)
        text = re.sub(r'\\end\{itemize\}', '', text)
        text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)  # Extract bold text
        text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)  # Extract italic text
        text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)    # Extract text
        text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)  # Remove other commands
        text = re.sub(r'\\([a-zA-Z]+)', r'\1', text)  # Remove backslash commands
        text = re.sub(r'\\\\', '', text)  # Remove line breaks
        text = re.sub(r'\\hline', '', text)  # Remove hline
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        return text.strip()
    
    def compile_tikz_to_image(self, tikz_code, filename_base):
        """Compile TikZ code to image"""
        # Create complete LaTeX document
        latex_doc = f"""
\\documentclass[border=5pt]{{standalone}}
\\usepackage{{tikz}}
\\usepackage{{amsmath}}
\\usepackage{{amssymb}}
\\begin{{document}}
{tikz_code}
\\end{{document}}
"""
        
        # Write to temporary file
        tex_file = os.path.join(self.temp_dir, f"{filename_base}.tex")
        with open(tex_file, 'w', encoding='utf-8') as f:
            f.write(latex_doc)
        
        try:
            # Compile with pdflatex
            subprocess.run(['pdflatex', '-interaction=nonstopmode', '-output-directory', self.temp_dir, tex_file], 
                         capture_output=True, check=True)
            
            # Convert PDF to PNG
            pdf_file = os.path.join(self.temp_dir, f"{filename_base}.pdf")
            png_file = os.path.join(self.temp_dir, f"{filename_base}.png")
            
            # Use pdftoppm for conversion (more reliable than Ghostscript)
            subprocess.run(['pdftoppm', '-png', '-r', '300', '-singlefile', pdf_file, 
                          os.path.join(self.temp_dir, filename_base)], 
                         capture_output=True, check=True)
            
            # Check if image was created
            if os.path.exists(png_file):
                return png_file
            else:
                # Try alternative conversion with ImageMagick
                subprocess.run(['convert', '-density', '300', pdf_file, png_file], 
                             capture_output=True, check=True)
                return png_file if os.path.exists(png_file) else None
                
        except subprocess.CalledProcessError as e:
            st.error(f"Error compiling TikZ: {e}")
            return None
        except Exception as e:
            st.error(f"Unexpected error: {e}")
            return None
    
    def create_word_document(self, exercises):
        """Create Word document from parsed exercises"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Bài tập', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for idx, exercise in enumerate(exercises, 1):
            # Add question number and text
            question_para = doc.add_paragraph()
            question_para.add_run(f'Câu {idx}. ').bold = True
            question_para.add_run(self.clean_latex_text(exercise['question']))
            
            # Add choices
            for i, choice in enumerate(exercise['choices']):
                choice_para = doc.add_paragraph()
                choice_para.add_run(f'    {chr(65 + i)}. ')  # A, B, C, D
                choice_para.add_run(self.clean_latex_text(choice))
            
            # Add TikZ image if exists
            if exercise['tikz']:
                image_file = self.compile_tikz_to_image(exercise['tikz'], f'tikz_{idx}')
                if image_file and os.path.exists(image_file):
                    doc.add_picture(image_file, width=Inches(3))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add solution if exists
            if exercise['solution']:
                doc.add_paragraph()
                solution_para = doc.add_paragraph()
                solution_para.add_run('Lời giải: ').bold = True
                solution_para.add_run(self.clean_latex_text(exercise['solution']))
            
            # Add spacing between exercises
            doc.add_paragraph()
        
        return doc
    
    def cleanup(self):
        """Clean up temporary files"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

def main():
    st.set_page_config(page_title="LaTeX to Word Converter", page_icon="📝")
    
    st.title("🔄 LaTeX to Word Converter")
    st.markdown("Chuyển đổi bài tập LaTeX sang định dạng Word")
    
    # Create two columns
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📥 Input LaTeX")
        
        # Text area for LaTeX input
        latex_input = st.text_area(
            "Nhập code LaTeX của bạn:",
            height=400,
            value=r"""\begin{ex}
\immini{Dựa vào hình vẽ (Hình b), hãy chọn khẳng định đúng trong các khẳng định sau đây?
\choice
{ Điểm $M$ nằm giữa $2$ điểm $N$ và $P$}
{ Điểm $N$ nằm giữa $2$ điểm $M$ và $P$}
{ Điểm $P$ nằm giữa $2$ điểm $M$ và $N$}
{ Hai điểm $M$ và $P$ nằm cùng phía đối với điểm $N$}
}{\begin{tikzpicture}[scale=1]
\coordinate (M) at (0.5, 0);
\coordinate (N) at (2.5, 0);
\coordinate (P) at (4.5, 0);
\draw[thick] (0, 0) -- (5.5, 0);
\foreach \pt/\angle in {M/90, N/90, P/90} {
\draw[fill=white] (\pt) circle (1.5pt) +(\angle:3mm) node{$\pt$};
}
\node[below=5mm of N] {Hình $b$};
\end{tikzpicture}
}
\loigiai{
Theo hình vẽ, các điểm $M$, $N$, $P$ nằm trên một đường thẳng theo thứ tự từ trái sang phải.
\begin{itemize}
\item Điểm $M$ nằm bên trái, điểm $P$ nằm bên phải
\item Điểm $N$ nằm giữa hai điểm $M$ và $P$
\item Hai điểm $M$ và $P$ nằm khác phía đối với điểm $N$
\end{itemize}
Vậy đáp án đúng là B.
}
\end{ex}"""
        )
        
        # File upload option
        uploaded_file = st.file_uploader("Hoặc tải lên file .tex", type=['tex'])
        if uploaded_file:
            latex_input = uploaded_file.read().decode('utf-8')
    
    with col2:
        st.subheader("📤 Output")
        
        if st.button("🔄 Chuyển đổi sang Word", type="primary"):
            if latex_input:
                try:
                    with st.spinner("Đang xử lý..."):
                        # Create converter instance
                        converter = LaTeXToWordConverter()
                        
                        # Extract and parse exercises
                        exercises_raw = converter.extract_exercises(latex_input)
                        exercises_parsed = [converter.parse_exercise(ex) for ex in exercises_raw]
                        
                        # Create Word document
                        doc = converter.create_word_document(exercises_parsed)
                        
                        # Save to BytesIO
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        # Clean up
                        converter.cleanup()
                        
                        # Offer download
                        st.success("✅ Chuyển đổi thành công!")
                        st.download_button(
                            label="📥 Tải xuống file Word",
                            data=doc_io.getvalue(),
                            file_name="exercises_converted.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        # Show preview
                        st.info(f"Đã chuyển đổi {len(exercises_parsed)} câu hỏi")
                        
                except Exception as e:
                    st.error(f"❌ Lỗi: {str(e)}")
            else:
                st.warning("⚠️ Vui lòng nhập nội dung LaTeX")
    
    # Instructions
    with st.expander("📖 Hướng dẫn sử dụng"):
        st.markdown("""
        ### Cấu trúc LaTeX được hỗ trợ:
        
        1. **Câu hỏi**: Đặt trong `\\begin{ex}...\\end{ex}`
        2. **Nội dung câu hỏi**: Trong `\\immini{...}` hoặc trước `\\choice`
        3. **Các lựa chọn**: Sau `\\choice`, mỗi lựa chọn trong `{...}`
        4. **Hình vẽ TikZ**: Trong `\\begin{tikzpicture}...\\end{tikzpicture}`
        5. **Lời giải**: Trong `\\loigiai{...}`
        
        ### Lưu ý:
        - Ứng dụng sẽ tự động biên dịch hình TikZ thành ảnh
        - Các ký hiệu toán học sẽ được chuyển đổi sang text thường
        - File Word xuất ra sẽ có format chuẩn với đánh số câu hỏi tự động
        """)

if __name__ == "__main__":
    main()
